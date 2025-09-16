# backend/app/services/content_service.py
"""
Content Management System for NCERT/CBSE curriculum content
"""

import json
import re
from typing import Dict, List, Optional, Tuple
import numpy as np
from datetime import datetime
import PyPDF2
import docx
import pandas as pd
from bs4 import BeautifulSoup
import requests
from sentence_transformers import SentenceTransformer
import faiss
import pickle

class ContentManagementService:
    """Manage educational content with RAG capabilities"""
    
    def __init__(self, config):
        self.config = config
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        self.vector_store = None
        self.content_index = {}
        self.initialize_vector_store()
        
    def initialize_vector_store(self):
        """Initialize FAISS vector store for content"""
        try:
            # Load existing index if available
            self.vector_store = faiss.read_index('content_vectors.index')
            with open('content_index.pkl', 'rb') as f:
                self.content_index = pickle.load(f)
        except:
            # Create new index
            dimension = 384  # Dimension of all-MiniLM-L6-v2
            self.vector_store = faiss.IndexFlatL2(dimension)
            self.content_index = {}
    
    async def import_ncert_content(self, file_path: str, subject: str, grade: int):
        """Import NCERT textbook content"""
        content_items = []
        
        if file_path.endswith('.pdf'):
            content_items = self._parse_pdf(file_path)
        elif file_path.endswith('.docx'):
            content_items = self._parse_docx(file_path)
        elif file_path.endswith('.html'):
            content_items = self._parse_html(file_path)
        
        # Process and structure content
        processed_content = []
        for item in content_items:
            processed = await self._process_content_item(item, subject, grade)
            processed_content.append(processed)
        
        # Generate embeddings and store
        await self._store_content_with_embeddings(processed_content)
        
        return len(processed_content)
    
    def _parse_pdf(self, file_path: str) -> List[Dict]:
        """Parse PDF content"""
        content_items = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            current_chapter = None
            current_section = None
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                # Detect chapter headers
                chapter_match = re.search(r'Chapter\s+(\d+)[:\s]+(.+)', text, re.IGNORECASE)
                if chapter_match:
                    current_chapter = {
                        'number': int(chapter_match.group(1)),
                        'title': chapter_match.group(2).strip(),
                        'page': page_num + 1
                    }
                
                # Detect sections
                section_match = re.search(r'(\d+\.\d+)\s+(.+)', text)
                if section_match:
                    current_section = {
                        'number': section_match.group(1),
                        'title': section_match.group(2).strip()
                    }
                
                # Extract content blocks
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    if len(para.strip()) > 50:  # Minimum content length
                        content_items.append({
                            'type': 'text',
                            'content': para.strip(),
                            'chapter': current_chapter,
                            'section': current_section,
                            'page': page_num + 1
                        })
                
                # Extract formulas
                formulas = re.findall(r'\$(.+?)\$', text)
                for formula in formulas:
                    content_items.append({
                        'type': 'formula',
                        'content': formula,
                        'chapter': current_chapter,
                        'section': current_section,
                        'page': page_num + 1
                    })
        
        return content_items
    
    def _parse_docx(self, file_path: str) -> List[Dict]:
        """Parse DOCX content"""
        content_items = []
        doc = docx.Document(file_path)
        
        current_chapter = None
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check for headers based on style
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                
                if level == 1:  # Chapter
                    current_chapter = {'title': text}
                elif level == 2:  # Section
                    current_section = {'title': text}
                
                content_items.append({
                    'type': 'heading',
                    'level': level,
                    'content': text,
                    'chapter': current_chapter,
                    'section': current_section
                })
            else:
                # Regular paragraph
                content_items.append({
                    'type': 'text',
                    'content': text,
                    'chapter': current_chapter,
                    'section': current_section
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            content_items.append({
                'type': 'table',
                'content': table_data,
                'chapter': current_chapter,
                'section': current_section
            })
        
        return content_items
    
    async def _process_content_item(self, item: Dict, subject: str, grade: int) -> Dict:
        """Process and enrich content item"""
        from ..models import db, Subject, Chapter, Topic, Content
        
        # Find or create subject
        subject_obj = Subject.query.filter_by(
            name=subject,
            grade=grade
        ).first()
        
        if not subject_obj:
            subject_obj = Subject(
                name=subject,
                grade=grade,
                board='CBSE'  # Default
            )
            db.session.add(subject_obj)
            db.session.flush()
        
        # Find or create chapter
        chapter_obj = None
        if item.get('chapter'):
            chapter_obj = Chapter.query.filter_by(
                subject_id=subject_obj.id,
                name=item['chapter'].get('title', '')
            ).first()
            
            if not chapter_obj:
                chapter_obj = Chapter(
                    subject_id=subject_obj.id,
                    name=item['chapter'].get('title', ''),
                    chapter_number=item['chapter'].get('number', 1)
                )
                db.session.add(chapter_obj)
                db.session.flush()
        
        # Extract key concepts and topics
        concepts = self._extract_concepts(item['content'])
        
        # Create content record
        content = Content(
            topic_id=None,  # Will be linked later
            content_type=item['type'],
            title=item.get('section', {}).get('title', ''),
            body=item['content'] if isinstance(item['content'], str) else json.dumps(item['content']),
            source=f"NCERT Grade {grade} {subject}",
            metadata={
                'chapter': item.get('chapter'),
                'section': item.get('section'),
                'page': item.get('page'),
                'concepts': concepts
            }
        )
        
        db.session.add(content)
        db.session.commit()
        
        return {
            'id': content.id,
            'content': item['content'],
            'type': item['type'],
            'metadata': content.metadata
        }
    
    def _extract_concepts(self, text: str) -> List[str]:
        """Extract key concepts from text"""
        import spacy
        
        # Load spaCy model (ensure it's downloaded: python -m spacy download en_core_web_sm)
        nlp = spacy.load("en_core_web_sm")
        
        doc = nlp(text)
        
        concepts = []
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ in ['PERSON', 'ORG', 'GPE', 'PRODUCT']:
                concepts.append(ent.text)
        
        # Extract noun phrases
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) <= 3:  # Max 3 words
                concepts.append(chunk.text)
        
        # Extract important terms (based on TF-IDF or similar)
        important_terms = self._get_important_terms(text)
        concepts.extend(important_terms)
        
        return list(set(concepts))[:10]  # Return top 10 unique concepts
    
    def _get_important_terms(self, text: str) -> List[str]:
        """Extract important terms using TF-IDF"""
        from sklearn.feature_extraction.text import TfidfVectorizer
        
        # Simple TF-IDF extraction
        vectorizer = TfidfVectorizer(max_features=5, stop_words='english')
        
        try:
            tfidf_matrix = vectorizer.fit_transform([text])
            feature_names = vectorizer.get_feature_names_out()
            return feature_names.tolist()
        except:
            return []
    
    async def _store_content_with_embeddings(self, content_items: List[Dict]):
        """Generate embeddings and store in vector database"""
        texts = [item['content'] if isinstance(item['content'], str) 
                else json.dumps(item['content']) for item in content_items]
        
        # Generate embeddings
        embeddings = self.embedder.encode(texts)
        
        # Add to FAISS index
        for i, (embedding, item) in enumerate(zip(embeddings, content_items)):
            index = len(self.content_index)
            self.vector_store.add(np.array([embedding]))
            self.content_index[index] = item
        
        # Save index
        faiss.write_index(self.vector_store, 'content_vectors.index')
        with open('content_index.pkl', 'wb') as f:
            pickle.dump(self.content_index, f)
    
    async def search_content(self, query: str, k: int = 5) -> List[Dict]:
        """Search content using semantic similarity"""
        # Generate query embedding
        query_embedding = self.embedder.encode([query])
        
        # Search in FAISS
        distances, indices = self.vector_store.search(query_embedding, k)
        
        # Retrieve content
        results = []
        for idx, distance in zip(indices[0], distances[0]):
            if idx in self.content_index:
                content = self.content_index[idx]
                content['similarity_score'] = 1 / (1 + distance)  # Convert distance to similarity
                results.append(content)
        
        return results
    
    async def generate_study_notes(self, topic_id: int) -> Dict:
        """Generate comprehensive study notes for a topic"""
        from ..models import Topic, Content
        
        topic = Topic.query.get(topic_id)
        if not topic:
            raise ValueError("Topic not found")
        
        # Get all content for topic
        content_items = Content.query.filter_by(topic_id=topic_id).all()
        
        # Organize content by type
        notes = {
            'topic': topic.name,
            'overview': '',
            'key_concepts': [],
            'formulas': [],
            'examples': [],
            'practice_questions': [],
            'summary': ''
        }
        
        for item in content_items:
            if item.content_type == 'text':
                notes['overview'] += item.body + '\n\n'
            elif item.content_type == 'formula':
                notes['formulas'].append(item.body)
            elif item.content_type == 'example':
                notes['examples'].append({
                    'title': item.title,
                    'content': item.body
                })
        
        # Extract key concepts
        all_text = notes['overview']
        notes['key_concepts'] = self._extract_concepts(all_text)
        
        # Generate summary using AI
        notes['summary'] = await self._generate_summary(all_text)
        
        return notes
    
    async def _generate_summary(self, text: str) -> str:
        """Generate summary using AI"""
        import openai
        
        openai.api_key = self.config['OPENAI_API_KEY']
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful educational assistant. Summarize the following content in a clear, concise manner suitable for students."},
                {"role": "user", "content": f"Summarize this educational content:\n\n{text[:2000]}"}  # Limit to 2000 chars
            ],
            max_tokens=200,
            temperature=0.7
        )
        
        return response.choices[0].message.content


# backend/app/routes/admin_routes.py
"""
Admin dashboard and content management routes
"""

from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from functools import wraps
import pandas as pd
from datetime import datetime, timedelta
from io import BytesIO
import json

from ..models import db, User, UserRole, Subject, Chapter, Topic, Content, Question
from ..models import Quiz, QuizResult, StudentProgress, Subscription, AnalyticsEvent
from ..services.content_service import ContentManagementService
from ..services.analytics_service import AnalyticsService

admin_bp = Blueprint('admin', __name__, url_prefix='/api/admin')
content_service = ContentManagementService(current_app.config)
analytics_service = AnalyticsService()

def admin_required(f):
    """Decorator to check if user is admin"""
    @wraps(f)
    @jwt_required()
    def decorated_function(*args, **kwargs):
        user_id = get_jwt_identity()
        user = User.query.get(user_id)
        
        if not user or user.role != UserRole.ADMIN:
            return jsonify({'error': 'Admin access required'}), 403
        
        return f(*args, **kwargs)
    
    return decorated_function

@admin_bp.route('/dashboard', methods=['GET'])
@admin_required
def admin_dashboard():
    """Get admin dashboard data"""
    
    # User statistics
    user_stats = {
        'total_users': User.query.count(),
        'students': User.query.filter_by(role=UserRole.STUDENT).count(),
        'parents': User.query.filter_by(role=UserRole.PARENT).count(),
        'teachers': User.query.filter_by(role=UserRole.TEACHER).count(),
        'active_today': User.query.filter(
            User.last_login >= datetime.utcnow() - timedelta(days=1)
        ).count(),
        'new_this_week': User.query.filter(
            User.created_at >= datetime.utcnow() - timedelta(weeks=1)
        ).count()
    }
    
    # Content statistics
    content_stats = {
        'total_subjects': Subject.query.count(),
        'total_chapters': Chapter.query.count(),
        'total_topics': Topic.query.count(),
        'total_questions': Question.query.count(),
        'total_content_items': Content.query.count()
    }
    
    # Activity statistics
    activity_stats = {
        'quizzes_today': QuizResult.query.filter(
            QuizResult.completed_at >= datetime.utcnow() - timedelta(days=1)
        ).count(),
        'active_subscriptions': Subscription.query.filter_by(status='active').count(),
        'revenue_this_month': db.session.query(
            db.func.sum(Subscription.amount)
        ).filter(
            Subscription.started_at >= datetime.utcnow() - timedelta(days=30),
            Subscription.status == 'active'
        ).scalar() or 0
    }
    
    # Recent activities
    recent_events = AnalyticsEvent.query.order_by(
        AnalyticsEvent.timestamp.desc()
    ).limit(10).all()
    
    return jsonify({
        'user_stats': user_stats,
        'content_stats': content_stats,
        'activity_stats': activity_stats,
        'recent_events': [e.to_dict() for e in recent_events]
    }), 200

@admin_bp.route('/users', methods=['GET'])
@admin_required
def list_users():
    """List all users with filtering"""
    
    # Pagination
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    
    # Filtering
    role = request.args.get('role')
    search = request.args.get('search')
    
    query = User.query
    
    if role:
        query = query.filter_by(role=role)
    
    if search:
        query = query.filter(
            db.or_(
                User.name.ilike(f'%{search}%'),
                User.email.ilike(f'%{search}%')
            )
        )
    
    # Paginate
    users = query.paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'users': [u.to_dict() for u in users.items],
        'total': users.total,
        'pages': users.pages,
        'current_page': page
    }), 200

@admin_bp.route('/user/<int:user_id>', methods=['GET', 'PUT', 'DELETE'])
@admin_required
def manage_user(user_id):
    """Manage individual user"""
    
    user = User.query.get_or_404(user_id)
    
    if request.method == 'GET':
        # Get detailed user info
        progress = StudentProgress.query.filter_by(student_id=user_id).all()
        subscriptions = Subscription.query.filter_by(user_id=user_id).all()
        
        return jsonify({
            'user': user.to_dict(),
            'progress': [p.to_dict() for p in progress],
            'subscriptions': [s.to_dict() for s in subscriptions]
        }), 200
    
    elif request.method == 'PUT':
        # Update user
        data = request.json
        
        if 'email' in data:
            user.email = data['email']
        if 'name' in data:
            user.name = data['name']
        if 'is_active' in data:
            user.is_active = data['is_active']
        if 'role' in data and data['role'] in ['student', 'parent', 'teacher']:
            user.role = data['role']
        
        db.session.commit()
        
        return jsonify({'message': 'User updated', 'user': user.to_dict()}), 200
    
    elif request.method == 'DELETE':
        # Soft delete user
        user.is_active = False
        db.session.commit()
        
        return jsonify({'message': 'User deactivated'}), 200

@admin_bp.route('/content/import', methods=['POST'])
@admin_required
async def import_content():
    """Import educational content"""
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    subject = request.form.get('subject')
    grade = request.form.get('grade', type=int)
    
    if not subject or not grade:
        return jsonify({'error': 'Subject and grade required'}), 400
    
    # Save file temporarily
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
        file.save(tmp.name)
        
        # Import content
        count = await content_service.import_ncert_content(
            tmp.name, 
            subject, 
            grade
        )
    
    return jsonify({
        'message': f'Successfully imported {count} content items',
        'count': count
    }), 201

@admin_bp.route('/content/subjects', methods=['GET', 'POST'])
@admin_required
def manage_subjects():
    """Manage subjects"""
    
    if request.method == 'GET':
        subjects = Subject.query.all()
        return jsonify({'subjects': [s.to_dict() for s in subjects]}), 200
    
    elif request.method == 'POST':
        data = request.json
        
        subject = Subject(
            name=data['name'],
            code=data.get('code'),
            board=data['board'],
            grade=data['grade'],
            is_core=data.get('is_core', True),
            icon=data.get('icon'),
            color=data.get('color')
        )
        
        db.session.add(subject)
        db.session.commit()
        
        return jsonify({'subject': subject.to_dict()}), 201

@admin_bp.route('/content/questions', methods=['GET', 'POST'])
@admin_required
def manage_questions():
    """Manage question bank"""
    
    if request.method == 'GET':
        # Filtering
        topic_id = request.args.get('topic_id', type=int)
        difficulty = request.args.get('difficulty')
        question_type = request.args.get('type')
        
        query = Question.query
        
        if topic_id:
            query = query.filter_by(topic_id=topic_id)
        if difficulty:
            query = query.filter_by(difficulty=difficulty)
        if question_type:
            query = query.filter_by(question_type=question_type)
        
        questions = query.all()
        
        return jsonify({'questions': [q.to_dict() for q in questions]}), 200
    
    elif request.method == 'POST':
        data = request.json
        
        question = Question(
            topic_id=data['topic_id'],
            chapter_id=data.get('chapter_id'),
            question_type=data['question_type'],
            difficulty=data['difficulty'],
            question_text=data['question_text'],
            question_image=data.get('question_image'),
            options=data.get('options'),
            correct_answer=data['correct_answer'],
            solution=data.get('solution'),
            solution_steps=data.get('solution_steps'),
            hints=data.get('hints'),
            marks=data.get('marks', 1.0),
            time_limit=data.get('time_limit'),
            skill_tags=data.get('skill_tags'),
            bloom_level=data.get('bloom_level')
        )
        
        db.session.add(question)
        db.session.commit()
        
        return jsonify({'question': question.to_dict()}), 201

@admin_bp.route('/content/bulk-upload', methods=['POST'])
@admin_required
def bulk_upload_questions():
    """Bulk upload questions from CSV/Excel"""
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    # Read file based on extension
    if file.filename.endswith('.csv'):
        df = pd.read_csv(file)
    elif file.filename.endswith(('.xlsx', '.xls')):
        df = pd.read_excel(file)
    else:
        return jsonify({'error': 'Invalid file format. Use CSV or Excel'}), 400
    
    # Validate columns
    required_columns = ['topic_id', 'question_text', 'question_type', 'correct_answer']
    missing_columns = [col for col in required_columns if col not in df.columns]
    
    if missing_columns:
        return jsonify({'error': f'Missing columns: {missing_columns}'}), 400
    
    # Process and insert questions
    questions_added = 0
    errors = []
    
    for index, row in df.iterrows():
        try:
            question = Question(
                topic_id=int(row['topic_id']),
                question_text=row['question_text'],
                question_type=row['question_type'],
                correct_answer=row['correct_answer'],
                difficulty=row.get('difficulty', 'medium'),
                options=json.loads(row['options']) if 'options' in row and pd.notna(row['options']) else None,
                solution=row.get('solution'),
                marks=float(row.get('marks', 1.0)),
                skill_tags=json.loads(row['skill_tags']) if 'skill_tags' in row and pd.notna(row['skill_tags']) else None
            )
            
            db.session.add(question)
            questions_added += 1
            
        except Exception as e:
            errors.append(f"Row {index + 2}: {str(e)}")
    
    db.session.commit()
    
    return jsonify({
        'message': f'Added {questions_added} questions',
        'questions_added': questions_added,
        'errors': errors
    }), 201

@admin_bp.route('/analytics/overview', methods=['GET'])
@admin_required
def analytics_overview():
    """Get platform-wide analytics"""
    
    # Time range
    days = request.args.get('days', 30, type=int)
    start_date = datetime.utcnow() - timedelta(days=days)
    
    # User growth
    user_growth = db.session.query(
        db.func.date(User.created_at).label('date'),
        db.func.count(User.id).label('count')
    ).filter(
        User.created_at >= start_date
    ).group_by(
        db.func.date(User.created_at)
    ).all()
    
    # Revenue growth
    revenue_growth = db.session.query(
        db.func.date(Subscription.started_at).label('date'),
        db.func.sum(Subscription.amount).label('revenue')
    ).filter(
        Subscription.started_at >= start_date,
        Subscription.status == 'active'
    ).group_by(
        db.func.date(Subscription.started_at)
    ).all()
    
    # Usage statistics
    usage_stats = db.session.query(
        AnalyticsEvent.event_type,
        db.func.count(AnalyticsEvent.id).label('count')
    ).filter(
        AnalyticsEvent.timestamp >= start_date
    ).group_by(
        AnalyticsEvent.event_type
    ).all()
    
    # Performance metrics
    avg_quiz_score = db.session.query(
        db.func.avg(QuizResult.percentage)
    ).filter(
        QuizResult.completed_at >= start_date
    ).scalar()
    
    avg_study_time = db.session.query(
        db.func.avg(StudentProgress.total_study_time)
    ).scalar()
    
    return jsonify({
        'user_growth': [{'date': str(d), 'count': c} for d, c in user_growth],
        'revenue_growth': [{'date': str(d), 'revenue': float(r)} for d, r in revenue_growth],
        'usage_stats': [{'type': t, 'count': c} for t, c in usage_stats],
        'performance_metrics': {
            'avg_quiz_score': float(avg_quiz_score) if avg_quiz_score else 0,
            'avg_study_time': float(avg_study_time) if avg_study_time else 0
        }
    }), 200

@admin_bp.route('/export/<report_type>', methods=['GET'])
@admin_required
def export_data(report_type):
    """Export data in various formats"""
    
    format_type = request.args.get('format', 'csv')
    
    if report_type == 'users':
        data = User.query.all()
        df = pd.DataFrame([u.to_dict() for u in data])
    elif report_type == 'progress':
        data = StudentProgress.query.all()
        df = pd.DataFrame([p.to_dict() for p in data])
    elif report_type == 'subscriptions':
        data = Subscription.query.all()
        df = pd.DataFrame([s.to_dict() for s in data])
    else:
        return jsonify({'error': 'Invalid report type'}), 400
    
    # Generate file
    output = BytesIO()
    
    if format_type == 'csv':
        df.to_csv(output, index=False)
        mimetype = 'text/csv'
        extension = 'csv'
    elif format_type == 'excel':
        df.to_excel(output, index=False)
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        extension = 'xlsx'
    else:
        return jsonify({'error': 'Invalid format'}), 400
    
    output.seek(0)
    
    return send_file(
        output,
        mimetype=mimetype,
        as_attachment=True,
        download_name=f'{report_type}_{datetime.now().strftime("%Y%m%d")}.{extension}'
    )

@admin_bp.route('/settings', methods=['GET', 'PUT'])
@admin_required
def manage_settings():
    """Manage system settings"""
    from ..models import SystemSettings
    
    if request.method == 'GET':
        settings = SystemSettings.query.all()
        return jsonify({
            'settings': {s.key: s.value for s in settings}
        }), 200
    
    elif request.method == 'PUT':
        data = request.json
        
        for key, value in data.items():
            setting = SystemSettings.query.filter_by(key=key).first()
            
            if setting:
                setting.value = value
            else:
                setting = SystemSettings(
                    key=key,
                    value=value
                )
                db.session.add(setting)
        
        db.session.commit()
        
        return jsonify({'message': 'Settings updated'}), 200

@admin_bp.route('/notifications/broadcast', methods=['POST'])
@admin_required
def broadcast_notification():
    """Send broadcast notification to users"""
    from ..services.notification_service import NotificationService
    
    data = request.json
    notification_service = NotificationService(current_app.config)
    
    # Get target users
    target = data.get('target', 'all')
    
    if target == 'all':
        users = User.query.filter_by(is_active=True).all()
    elif target == 'students':
        users = User.query.filter_by(role=UserRole.STUDENT, is_active=True).all()
    elif target == 'parents':
        users = User.query.filter_by(role=UserRole.PARENT, is_active=True).all()
    elif target == 'teachers':
        users = User.query.filter_by(role=UserRole.TEACHER, is_active=True).all()
    else:
        return jsonify({'error': 'Invalid target'}), 400
    
    # Send notifications
    sent_count = 0
    for user in users:
        try:
            notification_service.send_notification(
                user.id,
                data['type'],
                data['message']
            )
            
            if data.get('send_email'):
                notification_service.send_email(
                    user.email,
                    data.get('subject', 'ApexLearn Notification'),
                    data['message']
                )
            
            sent_count += 1
        except Exception as e:
            logger.error(f"Failed to send notification to user {user.id}: {e}")
    
    return jsonify({
        'message': f'Notification sent to {sent_count} users',
        'sent_count': sent_count,
        'total_users': len(users)
    }), 200